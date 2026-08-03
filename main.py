import json
import io
import os
import threading
from http.server import BaseHTTPRequestHandler, HTTPServer
from telegram import Update, KeyboardButton, ReplyKeyboardMarkup, WebAppInfo
from telegram.ext import ApplicationBuilder, ContextTypes, MessageHandler, CommandHandler, filters
from docx import Document

# --- RENDER İÇİN SAHTE WEB SUNUCUSU ---
class DummyHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.end_headers()
        self.wfile.write(b"Bot Aktif ve Calisiyor!")

def keep_alive():
    port = int(os.environ.get("PORT", 10000))
    server = HTTPServer(("0.0.0.0", port), DummyHandler)
    server.serve_forever()
# ----------------------------------------------------------------

BOT_TOKEN = os.environ.get("BOT_TOKEN")
NETLIFY_URL = "https://gleaming-florentine-2a9135.netlify.app"

async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    # Klavye butonlarımızı emojilerle renklendirip, evrak tiplerini URL parametresi olarak ekliyoruz.
    keyboard = [
        [
            KeyboardButton("👥 Veli Toplantısı", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=veli_toplanti")),
            KeyboardButton("📝 Sene Başı Zümre", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=sene_basi_zumre"))
        ],
        [
            KeyboardButton("🗳️ Başkanlık Seçimi", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=sinif_baskani")),
            KeyboardButton("🪑 Oturma Planı", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=oturma_plani"))
        ],
        [
            KeyboardButton("🎯 BEP Taslağı", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=bep")),
            KeyboardButton("🧭 Rehberlik Sevk", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=ogrenci_sevk"))
        ],
        [
            KeyboardButton("✂️ Ders Kesim Raporu", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=ders_kesim")),
            KeyboardButton("🤝 Veli Görüşme", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=veli_gorusme"))
        ],
        [
            KeyboardButton("✉️ Veli İzin Dilekçesi", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=veli_izin")),
            KeyboardButton("📅 Yıllık Ders Planı", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=yillik_plan"))
        ],
        [
            KeyboardButton("📖 Günlük Ders Planı", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=gunluk_plan")),
            KeyboardButton("⚖️ ŞÖK Tutanağı", web_app=WebAppInfo(url=f"{NETLIFY_URL}?evrak=sok"))
        ]
    ]
    
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)
    
    await update.message.reply_text(
        "Merhaba Öğretmenim! 👋\n\nAşağıdaki menüden hazırlamak istediğiniz evrağı seçiniz:",
        reply_markup=reply_markup
    )

async def web_app_data_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    raw_data = update.message.web_app_data.data
    data = json.loads(raw_data)
    
    evrak_turu = data.get("evrakTuru", "diger")
    okul_adi = data.get("okulAdi", "Belirtilmedi").upper()
    sinif = data.get("sinif", "Belirtilmedi")
    ogretmen_adi = data.get("ogretmenAdi", "Belirtilmedi")
    detaylar = data.get("detaylar", "")

    doc = Document()
    
    if evrak_turu == "veli_toplanti":
        dosya_adi = f"Veli_Toplanti_Tutanagi_{sinif}.docx"
        doc.add_heading(f"{okul_adi}\n{sinif} SINIFI VELİ TOPLANTI TUTANAĞI", level=1)
        doc.add_paragraph(f"Sınıf Rehber Öğretmeni: {ogretmen_adi}")
        doc.add_heading("Gündem ve Alınan Kararlar", level=2)
        doc.add_paragraph(detaylar if detaylar else "Toplantı gündem maddeleri görüşülmüş ve gerekli kararlar alınmıştır.")
        
    elif evrak_turu == "sene_basi_zumre":
        dosya_adi = "Sene_Basi_Zumre_Tutanagi.docx"
        doc.add_heading(f"{okul_adi}\nSENE BAŞI ZÜMRE ÖĞRETMENLER KURULU TUTANAĞI", level=1)
        doc.add_paragraph(f"Zümre Öğretmeni: {ogretmen_adi}")
        doc.add_heading("Gündem ve Alınan Kararlar", level=2)
        doc.add_paragraph(detaylar if detaylar else "Eğitim öğretim yılı planlaması yapılmış, müfredat değerlendirilmiştir.")
        
    elif evrak_turu == "sinif_baskani":
        dosya_adi = f"Sinif_Baskani_Secimi_{sinif}.docx"
        doc.add_heading(f"{okul_adi}\n{sinif} SINIFI BAŞKANLIK SEÇİMİ TUTANAĞI", level=1)
        doc.add_paragraph(f"Sınıf Rehber Öğretmeni: {ogretmen_adi}")
        doc.add_heading("Seçim Sonuçları ve Detaylar", level=2)
        doc.add_paragraph(detaylar if detaylar else "Sınıf başkanı ve yardımcısı demokratik oylama ile seçilmiştir.")

    elif evrak_turu == "oturma_plani":
        dosya_adi = f"Oturma_Plani_{sinif}.docx"
        doc.add_heading(f"{okul_adi}\n{sinif} SINIFI OTURMA PLANI", level=1)
        doc.add_paragraph(f"Sınıf Rehber Öğretmeni: {ogretmen_adi}")
        doc.add_heading("Plan Detayları", level=2)
        doc.add_paragraph(detaylar if detaylar else "Öğrencilerin fiziksel özellikleri ve pedagojik durumları göz önüne alınarak oturma planı oluşturulmuştur.")

    elif evrak_turu == "bep":
        dosya_adi = f"BEP_Plani_{sinif}.docx"
        doc.add_heading(f"{okul_adi}\nBİREYSELLEŞTİRİLMİŞ EĞİTİM PLANI (BEP)", level=1)
        doc.add_paragraph(f"Sorumlu Öğretmen: {ogretmen_adi} | Sınıf: {sinif}")
        doc.add_heading("Eğitsel Amaçlar ve Performans Notları", level=2)
        doc.add_paragraph(detaylar if detaylar else "Öğrencinin eğitsel performansı doğrultusunda amaçlar belirlenmiştir.")

    elif evrak_turu == "ogrenci_sevk":
        dosya_adi = f"Rehberlik_Sevk_{sinif}.docx"
        doc.add_heading(f"{okul_adi} REHBERLİK SERVİSİNE", level=1)
        doc.add_paragraph(f"Yönlendiren Öğretmen: {ogretmen_adi} | Sınıf: {sinif}")
        doc.add_heading("Gözlem ve Yönlendirme Nedeni", level=2)
        doc.add_paragraph(detaylar if detaylar else "Öğrencinin akademik/davranışsal gelişimi açısından rehberlik servisi ile görüşmesi uygun görülmüştür.")

    elif evrak_turu == "ders_kesim":
        dosya_adi = f"Ders_Kesim_Raporu.docx"
        doc.add_heading(f"{okul_adi}\nDERS KESİM RAPORU", level=1)
        doc.add_paragraph(f"Ders Öğretmeni: {ogretmen_adi} | Sınıf: {sinif}")
        doc.add_heading("Müfredat Gerçekleşme Durumu", level=2)
        doc.add_paragraph(detaylar if detaylar else "Yıllık planda belirtilen tüm konular işlenmiş ve ders kesimi yapılmıştır.")

    elif evrak_turu == "veli_gorusme":
        dosya_adi = f"Veli_Gorusme_Tutanagi_{sinif}.docx"
        doc.add_heading(f"{okul_adi}\nVELİ GÖRÜŞME TUTANAĞI", level=1)
        doc.add_paragraph(f"Görüşen Öğretmen: {ogretmen_adi} | Sınıf: {sinif}")
        doc.add_heading("Görüşme İçeriği", level=2)
        doc.add_paragraph(detaylar if detaylar else "Öğrencinin genel durumu hakkında veli ile görüşülmüş ve bilgilendirme yapılmıştır.")

    elif evrak_turu == "veli_izin":
        dosya_adi = f"Veli_Izin_Dilekcesi_{sinif}.docx"
        doc.add_heading("OKUL MÜDÜRLÜĞÜNE", level=1)
        doc.add_paragraph(f"Okul: {okul_adi} | Sınıf: {sinif}")
        doc.add_heading("İzin Talebi", level=2)
        doc.add_paragraph(detaylar if detaylar else "İlgili faaliyet/durum için velisi bulunduğum öğrencinin izinli sayılmasını arz ederim.")

    elif evrak_turu == "yillik_plan":
        dosya_adi = f"Yillik_Plan_{sinif}.docx"
        doc.add_heading(f"{okul_adi}\nYILLIK DERS PLANI", level=1)
        doc.add_paragraph(f"Ders Öğretmeni: {ogretmen_adi} | Sınıf: {sinif}")
        doc.add_heading("Plan Detayları", level=2)
        doc.add_paragraph(detaylar if detaylar else "Eğitim öğretim yılı müfredat ve kazanımları doğrultusunda hazırlanmıştır.")

    elif evrak_turu == "gunluk_plan":
        dosya_adi = f"Gunluk_Plan_{sinif}.docx"
        doc.add_heading(f"{okul_adi}\nGÜNLÜK DERS PLANI", level=1)
        doc.add_paragraph(f"Ders Öğretmeni: {ogretmen_adi} | Sınıf: {sinif}")
        doc.add_heading("Kazanımlar ve İşleniş", level=2)
        doc.add_paragraph(detaylar if detaylar else "İlgili ders saati için kazanımlar, yöntem ve teknikler belirlenmiştir.")

    elif evrak_turu == "sok":
        dosya_adi = f"SOK_Tutanagi_{sinif}.docx"
        doc.add_heading(f"{okul_adi}\n{sinif} SINIFI ŞÖK TOPLANTI TUTANAĞI", level=1)
        doc.add_paragraph(f"Sınıf Rehber Öğretmeni: {ogretmen_adi}")
        doc.add_heading("Alınan Kararlar ve Değerlendirmeler", level=2)
        doc.add_paragraph(detaylar if detaylar else "Sınıfın genel durumu ve öğrenci başarıları değerlendirilmiştir.")

    else:
        dosya_adi = f"Evrak_{sinif}.docx"
        doc.add_heading(f"{okul_adi} - EĞİTİM DOKÜMANI", level=1)
        doc.add_paragraph(f"Öğretmen: {ogretmen_adi} | Sınıf: {sinif}")
        doc.add_paragraph(detaylar)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    
    await update.message.reply_text("⏳ Evrağınız Word formatında hazırlanıyor...")
    await update.message.reply_document(
        document=target_stream,
        filename=dosya_adi,
        caption=f"✅ {okul_adi} - {sinif} belgeniz hazırlandı."
    )

def main():
    threading.Thread(target=keep_alive, daemon=True).start()
    
    app = ApplicationBuilder().token(BOT_TOKEN).build()
    
    app.add_handler(CommandHandler("start", start_command))
    app.add_handler(MessageHandler(filters.StatusUpdate.WEB_APP_DATA, web_app_data_handler))
    
    print("Bot aktif, öğretmenlerin evrak talepleri bekleniyor...")
    app.run_polling()

if __name__ == "__main__":
    main()
